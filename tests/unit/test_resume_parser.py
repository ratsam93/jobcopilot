from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi import HTTPException
import pytest

from apps.backend.app.modules.career_vault.service import store


def test_resume_parser_builds_profile_from_text_fixture() -> None:
    text = Path("tests/fixtures/sample_resume_ai_consultant.txt").read_text(encoding="utf-8")
    result = store.parse_resume_text(text, filename="sample_resume_ai_consultant.txt")

    assert result.parse_status == "parsed"
    assert result.created_profile.full_name == "Sam Patel"
    assert "AI automation" in [skill.skill_name for skill in result.created_profile.skills]
    assert result.created_profile.approved_claims


def test_resume_parser_extracts_text_from_docx_bytes() -> None:
    document = Document()
    document.add_paragraph("Alex Rivera")
    document.add_paragraph("alex@example.com")
    document.add_paragraph("+1 555 222 3333")
    document.add_paragraph("Remote USA")
    document.add_paragraph("- Built Python FastAPI SQL workflows")
    buffer = BytesIO()
    document.save(buffer)

    result = store.create_or_update_profile_from_bytes(buffer.getvalue(), filename="alex_resume.docx")

    assert result.parse_status == "parsed"
    assert result.created_profile.full_name == "Alex Rivera"
    assert result.created_profile.primary_email == "alex@example.com"
    assert "Python" in [skill.skill_name for skill in result.created_profile.skills]


def test_resume_parser_extracts_text_from_pdf_bytes() -> None:
    pdf = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> /MediaBox [0 0 612 792] /Contents 5 0 R >>
endobj
4 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
5 0 obj
<< /Length 108 >>
stream
BT /F1 12 Tf 72 720 Td (Jordan Lee) Tj 0 -20 Td (jordan@example.com) Tj 0 -20 Td (- Built Python FastAPI SQL workflows) Tj ET
endstream
endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000241 00000 n 
0000000311 00000 n 
trailer
<< /Root 1 0 R /Size 6 >>
startxref
469
%%EOF
"""

    result = store.create_or_update_profile_from_bytes(pdf, filename="jordan_resume.pdf")

    assert result.parse_status == "parsed"
    assert "Jordan Lee" in result.extracted_text
    assert result.created_profile.primary_email == "jordan@example.com"


def test_resume_parser_rejects_unreadable_pdf() -> None:
    with pytest.raises(HTTPException) as exc:
        store.create_or_update_profile_from_bytes(b"%PDF-1.4\n%%EOF", filename="empty_resume.pdf")

    assert exc.value.status_code == 422
