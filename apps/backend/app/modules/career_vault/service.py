from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from uuid import UUID

from fastapi import HTTPException
from docx import Document
from pypdf import PdfReader

from apps.backend.app.repositories import CareerProfileRepository
from apps.backend.app.shared.contracts import ResumeParseResult

from .models import (
    Bullet,
    BulletInput,
    CareerProfile,
    Claim,
    ClaimInput,
    DoNotClaim,
    DoNotClaimInput,
    ProfileUpdateInput,
    ResumeSource,
    ResumeUploadResult,
    Skill,
    SkillInput,
    utc_now,
)

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt"}
SKILL_KEYWORDS = [
    "ai automation",
    "python",
    "fastapi",
    "sql",
    "postgres",
    "docker",
    "aws",
    "azure",
    "gcp",
    "react",
    "typescript",
    "pydantic",
    "llm",
    "machine learning",
]


class CareerVaultStore:
    def __init__(self) -> None:
        self.repo = CareerProfileRepository()

    def _persist_profile(self, profile: CareerProfile, upload_key: str | None = None) -> None:
        key = upload_key
        if key is None:
            source = profile.source_resumes[0] if profile.source_resumes else None
            filename = source.filename if source else "resume.txt"
            text = source.text if source else ""
            key = self._upload_key(filename, text)
        self.repo.upsert(profile, key)

    def _detect_duplicate(self, filename: str, text: str) -> UUID | None:
        key = self._upload_key(filename, text)
        profile = self.repo.find_by_upload_key(key)
        return profile.candidate_profile_id if profile else None

    def _remember_upload(self, filename: str, text: str, profile_id: UUID) -> None:
        profile = self.get_profile(profile_id)
        self._persist_profile(profile, self._upload_key(filename, text))

    def _upload_key(self, filename: str, text: str) -> str:
        return f"{filename.lower()}::{re.sub(r'\\s+', ' ', text.strip()).lower()}"

    def get_profile(self, candidate_profile_id: UUID) -> CareerProfile:
        profile = self.repo.get(candidate_profile_id)
        if profile is None:
            raise HTTPException(status_code=404, detail="Candidate profile not found")
        return profile

    def update_profile(self, candidate_profile_id: UUID, payload: ProfileUpdateInput) -> CareerProfile:
        profile = self.get_profile(candidate_profile_id)
        data = payload.model_dump(exclude_none=True)
        for key, value in data.items():
            setattr(profile, key, value)
        profile.version += 1
        profile.updated_at = utc_now()
        self._persist_profile(profile)
        return profile

    def add_skill(self, candidate_profile_id: UUID, payload: SkillInput) -> Skill:
        profile = self.get_profile(candidate_profile_id)
        for skill in profile.skills:
            if (
                skill.skill_name.lower() == payload.skill_name.lower()
                and (skill.evidence_text or "").strip().lower() == (payload.evidence_text or "").strip().lower()
            ):
                return skill
        skill = Skill(candidate_profile_id=candidate_profile_id, **payload.model_dump())
        profile.skills.append(skill)
        profile.updated_at = utc_now()
        self._persist_profile(profile)
        return skill

    def add_claim(self, candidate_profile_id: UUID, payload: ClaimInput) -> Claim:
        profile = self.get_profile(candidate_profile_id)
        normalized_text = re.sub(r"\s+", " ", payload.claim_text).strip().lower()
        for claim in profile.approved_claims:
            if re.sub(r"\s+", " ", claim.claim_text).strip().lower() == normalized_text:
                return claim
        claim = Claim(candidate_profile_id=candidate_profile_id, **payload.model_dump())
        profile.approved_claims.append(claim)
        profile.updated_at = utc_now()
        self._persist_profile(profile)
        return claim

    def add_do_not_claim(self, candidate_profile_id: UUID, payload: DoNotClaimInput) -> DoNotClaim:
        profile = self.get_profile(candidate_profile_id)
        normalized_blocked = re.sub(r"\s+", " ", payload.blocked_claim).strip().lower()
        for entry in profile.do_not_claim:
            if re.sub(r"\s+", " ", entry.blocked_claim).strip().lower() == normalized_blocked:
                return entry
        entry = DoNotClaim(candidate_profile_id=candidate_profile_id, **payload.model_dump())
        profile.do_not_claim.append(entry)
        profile.updated_at = utc_now()
        self._persist_profile(profile)
        return entry

    def add_bullet(self, candidate_profile_id: UUID, payload: BulletInput) -> Bullet:
        profile = self.get_profile(candidate_profile_id)
        bullet = Bullet(candidate_profile_id=candidate_profile_id, **payload.model_dump())
        profile.bullet_bank.append(bullet)
        profile.updated_at = utc_now()
        self._persist_profile(profile)
        return bullet

    def parse_resume_text(self, text: str, filename: str = "resume.txt") -> ResumeUploadResult:
        if not text.strip():
            raise HTTPException(status_code=422, detail="Empty resume text requires manual input")
        existing = self._detect_duplicate(filename, text)
        if existing:
            profile = self.get_profile(existing)
            resume = ResumeSource(filename=filename, text=text.strip())
            if not any(
                source.filename.lower() == filename.lower()
                and re.sub(r"\s+", " ", source.text.strip()).lower() == re.sub(r"\s+", " ", text.strip()).lower()
                for source in profile.source_resumes
            ):
                profile.source_resumes.append(resume)
                profile.updated_at = utc_now()
                self._persist_profile(profile, self._upload_key(filename, text))
            result = ResumeUploadResult(
                candidate_profile_id=profile.candidate_profile_id,
                parse_status="parsed",
                source_filename=filename,
                extracted_text=text.strip(),
                created_profile=profile,
                duplicate_of=profile.candidate_profile_id,
            )
            ResumeParseResult.model_validate(result.model_dump(mode="json"))
            return result
        name = self._extract_name(text)
        email = self._extract_email(text)
        phone = self._extract_phone(text)
        location = self._extract_location(text)
        roles = self._infer_roles(text)
        profile = CareerProfile(
            full_name=name,
            primary_email=email,
            phone=phone,
            location=location,
            target_roles=roles,
            role_track=roles[0] if roles else "auto_detect",
            career_story=self._infer_story(text),
            source_resumes=[ResumeSource(filename=filename, text=text.strip())],
        )
        profile.skills = [
            Skill(candidate_profile_id=profile.candidate_profile_id, skill_name=skill, evidence_text=skill, approved=True)
            for skill in self._extract_skills(text)
        ]
        profile.approved_claims = [
            Claim(candidate_profile_id=profile.candidate_profile_id, claim_text=claim)
            for claim in self._extract_claims(text)
        ]
        self._persist_profile(profile, self._upload_key(filename, text))
        result = ResumeUploadResult(
            candidate_profile_id=profile.candidate_profile_id,
            parse_status="parsed",
            source_filename=filename,
            extracted_text=text.strip(),
            created_profile=profile,
        )
        ResumeParseResult.model_validate(result.model_dump(mode="json"))
        return result

    def create_or_update_profile_from_bytes(self, content: bytes, filename: str) -> ResumeUploadResult:
        suffix = Path(filename).suffix.lower()
        if suffix and suffix not in SUPPORTED_SUFFIXES:
            raise HTTPException(status_code=400, detail="Unsupported file type")
        text = self._extract_text_from_bytes(content, filename)
        if not text.strip():
            raise HTTPException(
                status_code=422,
                detail={
                    "message": "Could not extract readable resume text from uploaded file",
                    "filename": filename,
                    "supported_types": sorted(SUPPORTED_SUFFIXES),
                },
            )
        return self.parse_resume_text(text=text, filename=filename)

    def _extract_text_from_bytes(self, content: bytes, filename: str) -> str:
        suffix = Path(filename).suffix.lower()
        if suffix == ".pdf":
            try:
                reader = PdfReader(BytesIO(content))
                return "\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()
            except Exception as exc:
                raise HTTPException(status_code=422, detail=f"PDF text extraction failed: {exc}") from exc
        if suffix == ".docx":
            try:
                document = Document(BytesIO(content))
                paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
                table_cells: list[str] = []
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                table_cells.append(cell.text.strip())
                return "\n".join(paragraphs + table_cells).strip()
            except Exception as exc:
                raise HTTPException(status_code=422, detail=f"DOCX text extraction failed: {exc}") from exc
        return content.decode("utf-8-sig", errors="ignore")

    def _extract_email(self, text: str) -> str | None:
        matches = re.findall(r"[\w.\-+]+@[\w.\-]+\.\w+", text)
        if not matches:
            return None
        email = matches[0]
        local, domain = email.split("@", 1)
        compact_name = re.sub(r"[^a-z]", "", self._extract_name(text) or "", flags=re.IGNORECASE).lower()
        if compact_name and local.lower().startswith(compact_name):
            local = local[len(compact_name) :]
        elif len(local) > 32:
            local = re.split(r"(?<=[a-z])(?=[A-Z])", local)[-1]
        return f"{local}@{domain}" if local else email

    def _extract_phone(self, text: str) -> str | None:
        match = re.search(r"(\+?\d[\d\s().-]{7,}\d)", text)
        return match.group(0).strip() if match else None

    def _extract_name(self, text: str) -> str | None:
        first_line = next((line.strip() for line in text.splitlines() if line.strip()), "")
        return first_line if first_line and len(first_line.split()) <= 4 else None

    def _extract_location(self, text: str) -> str | None:
        for pattern in ("india", "remote", "united states", "usa"):
            if pattern in text.lower():
                return pattern.title()
        return None

    def _extract_skills(self, text: str) -> list[str]:
        lowered = text.lower()
        extracted: list[str] = []
        for skill in SKILL_KEYWORDS:
            if skill in lowered:
                if skill == "sql":
                    extracted.append("SQL")
                elif skill == "ai automation":
                    extracted.append("AI automation")
                elif skill == "llm":
                    extracted.append("LLM")
                else:
                    extracted.append(skill.title())
        return extracted

    def _infer_roles(self, text: str) -> list[str]:
        lowered = text.lower()
        roles: list[str] = []
        if any(word in lowered for word in ("backend", "fastapi", "api")):
            roles.append("backend engineer")
        if any(word in lowered for word in ("frontend", "react", "typescript")):
            roles.append("frontend engineer")
        if any(word in lowered for word in ("ml", "machine learning", "llm")):
            roles.append("ml engineer")
        return roles

    def _extract_claims(self, text: str) -> list[str]:
        claims: list[str] = []
        for line in text.splitlines():
            line = line.strip()
            if line.startswith(("-", "•", "*")) and len(line) > 8:
                claims.append(line.lstrip("-•* ").strip())
        return claims[:10]

    def _infer_story(self, text: str) -> str | None:
        claims = self._extract_claims(text)
        if not claims:
            return None
        return claims[0]

    def profile_summary(self, candidate_profile_id: UUID) -> dict[str, object]:
        profile = self.get_profile(candidate_profile_id)
        return {
            "profile": profile,
            "skills": profile.skills,
            "approved_claims": [claim for claim in profile.approved_claims if claim.approved],
            "do_not_claim": profile.do_not_claim,
            "bullet_bank": profile.bullet_bank,
        }

    def parsed_profile(self, candidate_profile_id: UUID) -> dict[str, object]:
        profile = self.get_profile(candidate_profile_id)
        skills = [skill.skill_name for skill in profile.skills]
        claims = [claim.claim_text for claim in profile.approved_claims if claim.approved]
        warnings: list[str] = []
        missing_info: list[str] = []
        if not profile.full_name:
            missing_info.append("name")
        if not profile.primary_email:
            missing_info.append("email")
        if not profile.phone:
            missing_info.append("phone")
        if not profile.location:
            missing_info.append("location")
        if not claims:
            warnings.append("No approved claims were detected from the resume.")
        if not skills:
            warnings.append("No skills were detected from the resume.")
        return {
            "candidate_name": profile.full_name,
            "email": profile.primary_email,
            "phone": profile.phone,
            "location": profile.location,
            "summary": profile.career_story or "Parsed from resume content.",
            "skills": skills,
            "experience": claims[:5],
            "projects": [bullet.bullet_text for bullet in profile.bullet_bank[:5]],
            "education": [],
            "certifications": [],
            "approved_claims_boundary": claims,
            "missing_info": missing_info,
            "parser_warnings": warnings,
            "raw_json": profile.model_dump(mode="json"),
        }


store = CareerVaultStore()
